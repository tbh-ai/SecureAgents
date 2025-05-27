#!/usr/bin/env python3
"""
Create a Word document from the Enhanced Adaptive Security Student Guide.
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("Installing python-docx...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

def create_word_document():
    """Create a professional Word document from the student guide."""
    
    # Create a new document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('🎓 Enhanced Adaptive Security System', 0)
    title_para = title.paragraph_format
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Student Guide', level=1)
    subtitle_para = subtitle.paragraph_format
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a line break
    doc.add_paragraph()
    
    # Table of Contents
    doc.add_heading('📚 Table of Contents', level=1)
    toc_items = [
        "1. What Is It?",
        "2. Why Is This Revolutionary?", 
        "3. Step-by-Step: How It Works",
        "4. Real-World Example",
        "5. Key Innovations",
        "6. Practical Benefits",
        "7. Summary"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # Section 1: What Is It?
    doc.add_heading('🌟 What Is It?', level=1)
    
    doc.add_paragraph(
        "Think of our enhanced adaptive security system like a super-smart security guard "
        "for your AI applications. But instead of just following a rulebook, this guard:"
    )
    
    features = [
        "🧠 Learns from every threat it sees",
        "🎭 Remembers how different users behave", 
        "📈 Gets smarter over time",
        "🔄 Adapts to new types of attacks",
        "⚡ Works lightning-fast (sub-millisecond responses)"
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Magic Formula
    doc.add_heading('The Magic Formula:', level=2)
    doc.add_paragraph("Traditional Security = Fixed rules that never change")
    doc.add_paragraph("Our System = Learning + Adaptation + Context + Speed", style='Intense Quote')
    
    # Section 2: Why Revolutionary
    doc.add_heading('🚀 Why Is This Revolutionary?', level=1)
    
    doc.add_heading('Traditional Security Systems (The Old Way):', level=2)
    old_way = [
        "❌ Fixed rules that never change",
        "❌ Can't learn from new attacks",
        "❌ Treats all users the same",
        "❌ High false positives", 
        "❌ Misses novel attacks",
        "❌ Slow and resource-heavy"
    ]
    
    for item in old_way:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Our Enhanced Adaptive System (The New Way):', level=2)
    new_way = [
        "✅ Learns and evolves continuously",
        "✅ Adapts to new attack patterns",
        "✅ Personalizes security per user",
        "✅ Reduces false positives over time",
        "✅ Detects novel attacks through behavioral analysis",
        "✅ Lightning-fast (sub-1ms) responses"
    ]
    
    for item in new_way:
        doc.add_paragraph(item, style='List Bullet')
    
    # Section 3: Step-by-Step
    doc.add_heading('📚 Step-by-Step: How It Works', level=1)
    
    # Step 1
    doc.add_heading('STEP 1: 📥 Input Reception', level=2)
    doc.add_paragraph("What happens when someone sends input:")
    
    # Add code example
    code_para = doc.add_paragraph()
    code_run = code_para.add_run('User Input: "import __builtins__; exec(__builtins__.__dict__[\'eval\'](\'malicious_code\'))"')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    
    doc.add_paragraph("System Response:")
    responses = [
        "🔍 Identifies the user: \"Oh, this is developer_123\"",
        "📋 Checks context: \"They're sending Python code at standard security level\"",
        "⏱️ Starts performance timer: Track how fast we can validate"
    ]
    
    for response in responses:
        doc.add_paragraph(response, style='List Bullet')
    
    doc.add_paragraph(
        "Think of it like: A bouncer at a club checking your ID and remembering if you're a regular customer.",
        style='Intense Quote'
    )
    
    # Step 2
    doc.add_heading('STEP 2: 🎭 Behavioral Analysis (Am I Acting Normal?)', level=2)
    doc.add_paragraph("The system analyzes user behavior:")
    
    # Add behavioral analysis example
    behavior_code = '''user_profile = {
    "typical_content": ["python_code", "javascript"],
    "normal_keywords": ["function", "import", "class"],
    "risk_score": 0.2,  # Low risk user
    "request_frequency": "normal"
}

current_behavior = {
    "content_type": "python_code",  # ✅ Normal for this user
    "keywords": ["import", "__builtins__", "exec", "eval"],  # ⚠️ Suspicious!
    "complexity": "high"  # ⚠️ More complex than usual
}

anomaly_score = 0.7  # High anomaly = suspicious behavior'''
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(behavior_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    doc.add_paragraph("What this means:")
    meanings = [
        "Normal behavior: User usually writes simple Python functions",
        "Current behavior: User is using advanced/dangerous Python features", 
        "Anomaly score: 0.7 out of 1.0 = \"This is unusual for this user!\""
    ]
    
    for meaning in meanings:
        doc.add_paragraph(meaning, style='List Bullet')
    
    doc.add_paragraph(
        "Think of it like: Your mom noticing you're acting weird - she knows your normal behavior!",
        style='Intense Quote'
    )
    
    # Step 3
    doc.add_heading('STEP 3: 🔍 Pattern Matching (Have I Seen This Attack Before?)', level=2)
    doc.add_paragraph("System checks against known attack patterns:")
    
    pattern_code = '''enhanced_patterns = [
    {
        "pattern": r"__builtins__.*eval",
        "category": "command_injection",
        "confidence": 0.95,
        "frequency": 15,  # Seen this 15 times before
        "last_seen": "2024-01-15"
    }
]

# Pattern matching result:
match_found = True
threat_confidence = 0.95
threat_category = "command_injection"'''
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(pattern_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    doc.add_paragraph("What happens:")
    happenings = [
        "System looks through its \"memory\" of attack patterns",
        "Finds a match: \"I've seen this `__builtins__` + `eval` combination before!\"",
        "Confidence: 95% sure this is a command injection attack",
        "Experience: \"I've caught this type of attack 15 times already\""
    ]
    
    for happening in happenings:
        doc.add_paragraph(happening, style='List Bullet')
    
    doc.add_paragraph(
        "Think of it like: A doctor recognizing symptoms they've seen many times before.",
        style='Intense Quote'
    )
    
    # Step 4
    doc.add_heading('STEP 4: 🎯 Smart Decision Making (Context + Behavior + Patterns)', level=2)
    doc.add_paragraph("System combines all information:")
    
    decision_code = '''base_threshold = 0.8  # Standard security level threshold
anomaly_adjustment = 0.7 * 0.2 = 0.14  # Lower threshold due to suspicious behavior
adjusted_threshold = 0.8 - 0.14 = 0.66  # Now more strict!

pattern_confidence = 0.95
context_boost = 0.05  # Boost because user behavior is suspicious
final_confidence = 0.95 + 0.05 = 1.0

# Decision:
if final_confidence (1.0) > adjusted_threshold (0.66):
    decision = "THREAT DETECTED!"'''
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(decision_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    doc.add_paragraph("What this means:")
    decision_meanings = [
        "Normal situation: Would need 80% confidence to block",
        "Suspicious user: Only need 66% confidence (more strict)",
        "Pattern confidence: 95% + 5% boost = 100% sure it's a threat",
        "Decision: BLOCK IT!"
    ]
    
    for meaning in decision_meanings:
        doc.add_paragraph(meaning, style='List Bullet')
    
    doc.add_paragraph(
        "Think of it like: Airport security being extra careful with someone acting suspiciously.",
        style='Intense Quote'
    )
    
    # Add page break before continuing
    doc.add_page_break()
    
    # Step 5
    doc.add_heading('STEP 5: 🧬 Learning & Evolution (Getting Smarter)', level=2)
    doc.add_paragraph("System learns from this detection:")
    
    learning_code = '''# Pattern gets stronger:
pattern.frequency += 1  # Now seen 16 times instead of 15
pattern.confidence = recalculate_confidence()  # Might increase to 0.96
pattern.last_seen = "today"

# User profile updates:
user_profile.risk_score += 0.1  # User becomes slightly more risky
user_profile.suspicious_keywords.add("__builtins__")

# Memory storage:
attack_history.append({
    "user": "developer_123",
    "attack_type": "command_injection", 
    "blocked": True,
    "timestamp": "now"
})'''
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(learning_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    doc.add_paragraph("What the system learns:")
    learnings = [
        "Pattern gets stronger: \"I'm even more confident about this attack type now\"",
        "User profile updates: \"This user tried something suspicious\"",
        "Memory storage: \"I'll remember this happened\""
    ]
    
    for learning in learnings:
        doc.add_paragraph(learning, style='List Bullet')
    
    doc.add_paragraph(
        "Think of it like: Your immune system getting stronger after fighting off a virus.",
        style='Intense Quote'
    )
    
    # Step 6
    doc.add_heading('STEP 6: 🔄 Hybrid Validation (Multiple Security Layers)', level=2)
    doc.add_paragraph("Our system uses THREE layers of protection:")
    
    hybrid_code = '''# Layer 1: Fast Regex Check (milliseconds)
regex_result = check_dangerous_patterns(text)

# Layer 2: Machine Learning (few milliseconds)  
ml_result = analyze_with_ai_model(text)

# Layer 3: Large Language Model (if needed)
llm_result = ask_smart_ai_to_analyze(text)

# Combine results:
final_decision = combine_all_results(regex_result, ml_result, llm_result)'''
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(hybrid_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    doc.add_paragraph("Why multiple layers?")
    layers = [
        "Regex: Super fast, catches obvious attacks",
        "ML: Catches subtle patterns, still fast",
        "LLM: Understands context and meaning, slower but very smart"
    ]
    
    for layer in layers:
        doc.add_paragraph(layer, style='List Bullet')
    
    doc.add_paragraph(
        "Think of it like: Airport security with metal detectors, X-ray machines, AND human guards.",
        style='Intense Quote'
    )
    
    # Step 7
    doc.add_heading('STEP 7: ⚡ Lightning-Fast Response', level=2)
    doc.add_paragraph("Total time breakdown:")
    
    speed_code = '''behavioral_analysis = 0.1ms
pattern_matching = 0.1ms  
decision_making = 0.05ms
learning_update = 0.05ms
total_time = 0.3ms  # Less than 1 millisecond!

response = {
    "is_secure": False,
    "confidence": 1.0,
    "threat_type": "command_injection",
    "reason": "Dangerous Python builtin manipulation detected",
    "suggestions": ["Remove __builtins__ access", "Use safer alternatives"],
    "time_taken": "0.3ms"
}'''
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(speed_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    doc.add_paragraph("Speed comparison:")
    speeds = [
        "Blinking your eye: ~300ms",
        "Our security check: 0.3ms",
        "We're 1000x faster than an eye blink!"
    ]
    
    for speed in speeds:
        doc.add_paragraph(speed, style='List Bullet')
    
    # Summary Section
    doc.add_page_break()
    doc.add_heading('🎓 Summary for Students', level=1)
    
    doc.add_paragraph("Think of our system as:")
    summary_points = [
        "🧠 A learning security guard that remembers every threat",
        "🎭 A behavioral analyst that knows how users normally act",
        "🔄 A team of specialists (regex, ML, LLM) working together",
        "⚡ A lightning-fast decision maker (sub-millisecond responses)",
        "🧬 An evolving organism that adapts to new threats"
    ]
    
    for point in summary_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph("The magic happens because:")
    magic_points = [
        "It learns from experience (like humans do)",
        "It considers context (who, what, when, where)",
        "It uses multiple perspectives (different AI techniques)",
        "It adapts in real-time (no waiting for updates)",
        "It personalizes security (different rules for different users)"
    ]
    
    for point in magic_points:
        doc.add_paragraph(point, style='List Number')
    
    # Final paragraph
    final_para = doc.add_paragraph(
        "This represents the future of AI security - systems that don't just follow rules, "
        "but actually understand, learn, and evolve! 🚀"
    )
    final_para.style = 'Intense Quote'
    
    # Key Takeaways
    doc.add_heading('🎯 Key Takeaways', level=1)
    takeaways = [
        "Adaptive Learning: The system gets smarter with every interaction",
        "Behavioral Analysis: It knows what's normal vs. suspicious for each user",
        "Multi-Layer Defense: Three different AI techniques working together",
        "Real-Time Performance: Sub-millisecond responses for production use",
        "Context Awareness: Understands the situation, not just the text",
        "Continuous Evolution: Patterns and confidence levels improve over time"
    ]
    
    for takeaway in takeaways:
        doc.add_paragraph(takeaway, style='List Number')
    
    # Final statement
    final_statement = doc.add_paragraph(
        "This is not just security - it's intelligent, adaptive, learning security that "
        "represents the cutting edge of AI protection technology! 🚀"
    )
    final_statement.style = 'Intense Quote'
    
    # Save the document
    doc.save('Enhanced_Adaptive_Security_Student_Guide.docx')
    print("✅ Word document created: Enhanced_Adaptive_Security_Student_Guide.docx")

if __name__ == "__main__":
    create_word_document()
